import io
import pandas as pd
from docx import Document
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document as LangDocument
from config import settings

def load_vector_db():
    df = pd.read_excel(settings.TAGS_FILE)
    df.columns = [str(c).strip() for c in df.columns]

    tag_docs = [
        LangDocument(page_content=f"Tag: {r['Tag']}...", metadata={"tag_name": r['Tag']}) 
        for _, r in df.iterrows()
    ]

    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

    # ADD THIS PARAMETER:
    return Chroma.from_documents(
        documents=tag_docs, 
        embedding=embeddings,
        collection_name="excel_tags",
        persist_directory="./chroma_db"  # <--- This creates the folder!
    )

def parse_docx_table(file_content):
    """Converts a Word table to Markdown using an In-Memory buffer."""
    # BytesIO wraps the binary data so 'python-docx' thinks it's a file
    doc = Document(io.BytesIO(file_content))
    
    if not doc.tables:
        return None

    table = doc.tables[0] # Assumes the first table contains the data
    md_output = ""
    
    for row in table.rows:
        text_row = [cell.text.strip() for cell in row.cells]
        md_output += "| " + " | ".join(text_row) + " |\n"
    
    return md_output

def get_tag_from_llm(fact, table_context, tag_db):
    """Queries OpenRouter using the candidate tags from ChromaDB."""
    # 1. Similarity Search (Find top 3 most likely tags)
    relevant_tags = tag_db.similarity_search(fact, k=3)
    tag_options = "\n\n".join([d.page_content for d in relevant_tags])

    # 2. Setup LLM with explicit OpenRouter credentials from config.py
    llm = ChatOpenAI(
        model=settings.MODEL_NAME, 
        openai_api_key=settings.OPENROUTER_KEY,
        openai_api_base=settings.BASE_URL,
        temperature=0.1
    )

    prompt = f"""
    Context Table:
    {table_context}
    
    Target Fact: "{fact}"
    
    Potential Tags:
    {tag_options}
    
    Task: Identify the correct Tag for the Target Fact based on the table.
    Return ONLY the Tag Name. No explanation.
    """
    
    try:
        response = llm.invoke(prompt)
        # Clean the output to ensure only the tag string is returned
        return response.content.strip().split('\n')[0].replace("Tag:", "").strip()
    except Exception as e:
        return f"Error calling LLM: {str(e)}"